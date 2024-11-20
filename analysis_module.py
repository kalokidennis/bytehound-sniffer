import os
import csv
from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QLabel, QTableWidget, QTableWidgetItem,
    QPushButton, QMessageBox, QFileDialog, QSizePolicy
)
from PyQt5.QtCore import Qt
from docx import Document


class PacketAnalysisWindow(QDialog):
    def __init__(self, packets, parent=None):
        super().__init__(parent)
        self.packets = packets
        self.setWindowTitle("Packet Analysis Window")
        self.setMinimumSize(600, 400)
        self.setWindowFlags(self.windowFlags() | Qt.WindowMinMaxButtonsHint)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)

        # Title label
        title_label = QLabel("Captured Packets Analysis")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)

        # Table to display captured packets
        self.packet_table = QTableWidget()
        self.packet_table.setColumnCount(9)
        self.packet_table.setHorizontalHeaderLabels([
            'Source MAC', 'Destination MAC', 'Source IP', 'Destination IP',
            'Source Port', 'Destination Port', 'Protocol', 'Summary', 'Safe'
        ])
        self.packet_table.setRowCount(len(self.packets))
        self.populate_table()
        self.packet_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.packet_table.horizontalHeader().setStretchLastSection(True)
        layout.addWidget(self.packet_table)

        # Analyze button
        analyze_button = QPushButton("Analyze Packets")
        analyze_button.clicked.connect(self.analyze_packets)
        layout.addWidget(analyze_button)

        # Close button
        close_button = QPushButton("Close")
        close_button.clicked.connect(self.close)
        layout.addWidget(close_button)

        self.setLayout(layout)

    def populate_table(self):
        """Populate the table with packet data."""
        for row_index, packet in enumerate(self.packets):
            self.packet_table.setItem(row_index, 0, QTableWidgetItem(packet.get('eth_src', 'Unknown')))
            self.packet_table.setItem(row_index, 1, QTableWidgetItem(packet.get('eth_dst', 'Unknown')))
            self.packet_table.setItem(row_index, 2, QTableWidgetItem(packet.get('ip_src', 'Unknown')))
            self.packet_table.setItem(row_index, 3, QTableWidgetItem(packet.get('ip_dst', 'Unknown')))
            self.packet_table.setItem(row_index, 4, QTableWidgetItem(str(packet.get('sport', '-'))))
            self.packet_table.setItem(row_index, 5, QTableWidgetItem(str(packet.get('dport', '-'))))
            self.packet_table.setItem(row_index, 6, QTableWidgetItem(packet.get('protocol', 'Unknown')))
            self.packet_table.setItem(row_index, 7, QTableWidgetItem(packet.get('summary', 'No summary available')))
            self.packet_table.setItem(row_index, 8, QTableWidgetItem('Yes' if packet.get('safe', True) else 'No'))

    def analyze_packets(self):
        """Analyze captured packets and show results."""
        try:
            if not self.packets:
                QMessageBox.warning(self, "Warning", "No packets to analyze!")
                return

            # Perform the analysis
            analysis_result, protocol_counts = self.perform_analysis()
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Analysis Result")
            msg_box.setText(analysis_result)
            msg_box.setStandardButtons(QMessageBox.Ok | QMessageBox.Save)

            save_button = msg_box.button(QMessageBox.Save)
            save_button.setText("Save Report")

            result = msg_box.exec_()
            if result == QMessageBox.Save:
                self.generate_report(analysis_result, protocol_counts)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred during analysis: {e}")

    def perform_analysis(self):
        """Perform analysis on captured packets and check for potential safety risks."""
        total_packets = len(self.packets)
        safe_packets = sum(1 for packet in self.packets if packet.get('safe', True))
        unsafe_packets = total_packets - safe_packets

        # Count packets per protocol
        protocol_counts = {}
        for packet in self.packets:
            protocol = packet.get('protocol', 'Unknown')
            protocol_counts[protocol] = protocol_counts.get(protocol, 0) + 1

        analysis_result = (
            f"Total Packets Captured: {total_packets}\n"
            f"Safe Packets: {safe_packets} ({(safe_packets / total_packets) * 100:.2f}%)\n"
            f"Unsafe Packets: {unsafe_packets} ({(unsafe_packets / total_packets) * 100:.2f}%)\n"
            f"\nPackets per Protocol:\n" +
            "\n".join([f"{protocol}: {count} packets" for protocol, count in protocol_counts.items()])
        )

        return analysis_result, protocol_counts

    def generate_report(self, analysis_result, protocol_counts):
        """Generate and save a report with packet data and analysis results."""
        save_path, selected_filter = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            os.path.expanduser("~"),
            "Word Files (*.docx);;CSV Files (*.csv);;Text Files (*.txt);;All Files (*)"
        )

        if not save_path:
            return

        try:
            if selected_filter == "Word Files (*.docx)" or save_path.endswith(".docx"):
                self.save_report_as_docx(save_path, analysis_result)
            elif selected_filter == "CSV Files (*.csv)" or save_path.endswith(".csv"):
                self.save_report_as_csv(save_path)
            else:
                self.save_report_as_txt(save_path, analysis_result)

            QMessageBox.information(self, "Success", f"Report saved to {save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not save report: {e}")

    def save_report_as_csv(self, save_path):
        """Save the report as a CSV file, mirroring the UI table structure."""
        with open(save_path, 'w', newline='') as csvfile:
            writer = csv.writer(csvfile)

            # Write table headers
            headers = [self.packet_table.horizontalHeaderItem(i).text() for i in range(self.packet_table.columnCount())]
            writer.writerow(headers)

            # Write table rows
            for row in range(self.packet_table.rowCount()):
                row_data = [self.packet_table.item(row, col).text() if self.packet_table.item(row, col) else '' for col in range(self.packet_table.columnCount())]
                writer.writerow(row_data)

    def save_report_as_txt(self, save_path, analysis_result):
        """Save the report as a TXT file, including the table as visible in the UI."""
        with open(save_path, 'w') as txtfile:
            txtfile.write("Packet Analysis Report\n")
            txtfile.write("=" * 50 + "\n")
            txtfile.write(f"{analysis_result}\n\n")
            txtfile.write("Captured Packet Details:\n")
            txtfile.write("=" * 50 + "\n")

            # Write table headers
            headers = [self.packet_table.horizontalHeaderItem(i).text() for i in range(self.packet_table.columnCount())]
            txtfile.write("\t".join(headers) + "\n")

            # Write table rows
            for row in range(self.packet_table.rowCount()):
                row_data = [self.packet_table.item(row, col).text() if self.packet_table.item(row, col) else '' for col in range(self.packet_table.columnCount())]
                txtfile.write("\t".join(row_data) + "\n")

    def save_report_as_docx(self, save_path, analysis_result):
        """Save the report as a DOCX file, with a well-formatted table."""
        doc = Document()

        # Add title
        doc.add_heading("Packet Analysis Report", level=1)
        doc.add_paragraph(analysis_result)

        # Add table
        table = doc.add_table(rows=self.packet_table.rowCount() + 1, cols=self.packet_table.columnCount())
        table.style = 'Table Grid'

        # Add headers
        headers = [self.packet_table.horizontalHeaderItem(i).text() for i in range(self.packet_table.columnCount())]
        for col, header in enumerate(headers):
            table.cell(0, col).text = header

        # Add rows
        for row in range(self.packet_table.rowCount()):
            for col in range(self.packet_table.columnCount()):
                table.cell(row + 1, col).text = self.packet_table.item(row, col).text() if self.packet_table.item(row, col) else ''

        doc.save(save_path)
